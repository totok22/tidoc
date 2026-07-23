"""打印导出组件测试：PDF 拼接、图片转 PDF、Word 生成、抬头强隔离。"""

import glob
import json
import subprocess
import sys
from decimal import Decimal

import pytest

import tidoc_print

# 组件依赖未装则整体跳过（核心测试不受影响）
pytestmark = pytest.mark.skipif(not tidoc_print.is_available(), reason="打印组件依赖未安装")

SAMPLE_DIR = "/Users/poli/invoice2docx/invoices"


def _sample(pattern):
    files = sorted(glob.glob(f"{SAMPLE_DIR}/{pattern}"))
    if not files:
        pytest.skip(f"无样本：{pattern}")
    return files


def test_availability():
    assert tidoc_print.is_available()
    assert tidoc_print.missing_dependencies() == []


def test_component_self_test_imports_full_print_stack():
    proc = subprocess.run(
        [sys.executable, "-m", "tidoc_print", "--self-test"],
        text=True,
        capture_output=True,
        check=False,
    )

    assert proc.returncode == 0, proc.stderr
    assert json.loads(proc.stdout)["ok"] is True


def test_merge_pdfs(tmp_path):
    from tidoc_print.pdf_merge import merge_pdfs
    from pypdf import PdfReader
    pdfs = _sample("*发票.pdf")[:2]
    out = merge_pdfs(pdfs, tmp_path / "merged.pdf")
    assert out.exists()
    assert len(PdfReader(str(out)).pages) >= 2


def test_merge_pdfs_with_annotation(tmp_path):
    from tidoc_print.pdf_merge import merge_pdfs
    pdfs = _sample("*发票.pdf")[:1]
    out = merge_pdfs(pdfs, tmp_path / "annotated.pdf", annotations=["发票号：123   报账人：张三"])
    assert out.exists() and out.stat().st_size > 0


def test_images_to_pdf(tmp_path):
    from tidoc_print.pdf_merge import images_to_pdf
    from pypdf import PdfReader
    imgs = _sample("*付款截图.jpg")[:2]
    out = images_to_pdf(imgs, tmp_path / "pay.pdf", annotations=["实付：¥100"] * len(imgs))
    assert out.exists()
    reader = PdfReader(str(out))
    assert len(reader.pages) == 1
    page = reader.pages[0]
    assert float(page.mediabox.width) > float(page.mediabox.height)


def test_images_to_pdf_two_per_page(tmp_path):
    from tidoc_print.pdf_merge import images_to_pdf
    from pypdf import PdfReader
    imgs = (_sample("*付款截图.jpg") * 2)[:3]
    out = images_to_pdf(imgs, tmp_path / "pay3.pdf", annotations=["No.1-1/3", "No.1-2/3", "No.1-3/3"])
    assert out.exists()
    reader = PdfReader(str(out))
    assert len(reader.pages) == 2
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "Page 1/2" in text


def test_generate_word_docs(tmp_path):
    from tidoc_print import PrintEntry, PrintItem, PersonProfile
    from tidoc_print.word_docs import generate_acceptance_doc, generate_reimburse_doc
    from docx import Document

    entries = [
        PrintEntry(
            entry_id="e1", title="北京理工大学", invoice_no="123",
            seller="某某公司", total=Decimal("100.00"), profile_name="张三",
            items=[PrintItem(actual_name="电阻", unit="个", quantity=Decimal("10"), total=Decimal("100.00"), seller="某某公司", invoice_no="123")],
        ),
    ]
    r = generate_reimburse_doc(entries, tmp_path / "报账说明.docx", "2026年7月5日", PersonProfile(person_name="张三"))
    a = generate_acceptance_doc(entries, tmp_path / "验收单.docx", "2026年7月5日")
    assert r.exists() and a.exists()
    # 报账说明表格里应有数据行
    assert len(Document(str(r)).tables[0].rows) >= 2
    assert len(Document(str(a)).tables[0].rows) >= 2


def test_title_isolation(tmp_path):
    """两个抬头必须分成两套文件，绝不混合（第 7 节）。"""
    from tidoc_print import PrintEntry, PrintItem, build_print_package, PrintOptions

    def mk(title, no):
        return PrintEntry(
            entry_id=no, title=title, invoice_no=no, seller="公司", total=Decimal("50.00"),
            profile_name="张三",
            items=[PrintItem(actual_name="物资", unit="个", quantity=Decimal("1"), total=Decimal("50.00"), invoice_no=no)],
        )

    entries = [mk("北京理工大学", "1"), mk("北京理工大学教育基金会", "2")]
    opts = PrintOptions(make_invoice_pdf=False, make_payment_pdf=False, make_inspection_pdf=False)
    results = build_print_package(entries, tmp_path, opts)
    titles = {r.title for r in results}
    assert titles == {"北京理工大学", "北京理工大学教育基金会"}
    # 两个抬头各自的目录独立
    assert (tmp_path / "北京理工大学").exists()
    assert (tmp_path / "北京理工大学教育基金会").exists()


def test_full_package_with_real_files(tmp_path):
    from tidoc_print import PrintEntry, PrintItem, build_print_package, PrintOptions
    pdf = _sample("*发票.pdf")[0]
    img = _sample("*付款截图.jpg")[0]
    insp = _sample("*查验单.pdf")[0]
    entry = PrintEntry(
        entry_id="e1", title="北京理工大学", invoice_no="123", seller="公司",
        total=Decimal("100.00"), paid_amount="100.00", profile_name="张三",
        items=[PrintItem(actual_name="物资", unit="个", quantity=Decimal("1"), total=Decimal("100.00"), invoice_no="123")],
        invoice_pdfs=[pdf], payment_images=[img], inspection_pdfs=[insp],
    )
    results = build_print_package([entry], tmp_path, PrintOptions())
    assert len(results) == 1
    files = results[0].files
    assert "entry_bundle_pdf" in files
    assert "invoice_pdf" not in files and "payment_pdf" not in files and "inspection_pdf" not in files
    assert "reimburse_doc" in files and "acceptance_doc" in files
    for path in files.values():
        from pathlib import Path
        assert Path(path).exists()


def test_separate_material_pdfs_remain_optional(tmp_path):
    from tidoc_print import PrintEntry, PrintOptions, build_print_package
    pdf = _sample("*发票.pdf")[0]
    img = _sample("*付款截图.jpg")[0]
    insp = _sample("*查验单.pdf")[0]
    entry = PrintEntry(
        entry_id="e1", title="北京理工大学", invoice_no="123",
        invoice_pdfs=[pdf], payment_images=[img], inspection_pdfs=[insp],
    )
    options = PrintOptions(
        make_entry_bundle_pdf=False,
        make_invoice_pdf=True,
        make_payment_pdf=True,
        make_inspection_pdf=True,
        make_reimburse_doc=False,
        make_acceptance_doc=False,
    )
    files = build_print_package([entry], tmp_path, options)[0].files
    assert set(files) == {"invoice_pdf", "payment_pdf", "inspection_pdf"}


def test_entry_bundle_keeps_each_entry_materials_together(tmp_path):
    from PIL import Image
    from pypdf import PdfReader
    from reportlab.pdfgen import canvas
    from tidoc_print import PrintEntry, PrintOptions, build_print_package

    def make_pdf(path, text):
        doc = canvas.Canvas(str(path))
        doc.drawString(72, 760, text)
        doc.save()
        return str(path)

    def make_image(path, color):
        Image.new("RGB", (320, 640), color).save(path)
        return str(path)

    entries = [
        PrintEntry(
            entry_id="a", title="北京理工大学",
            invoice_pdfs=[make_pdf(tmp_path / "invoice-a.pdf", "INVOICE-A")],
            payment_images=[make_image(tmp_path / "pay-a-1.png", "red"),
                            make_image(tmp_path / "pay-a-2.png", "blue")],
            inspection_pdfs=[make_pdf(tmp_path / "inspection-a.pdf", "INSPECTION-A")],
        ),
        PrintEntry(
            entry_id="b", title="北京理工大学",
            invoice_pdfs=[make_pdf(tmp_path / "invoice-b.pdf", "INVOICE-B")],
            payment_images=[make_image(tmp_path / "pay-b.png", "green")],
            inspection_pdfs=[make_pdf(tmp_path / "inspection-b.pdf", "INSPECTION-B")],
        ),
    ]
    options = PrintOptions(make_reimburse_doc=False, make_acceptance_doc=False)
    result = build_print_package(entries, tmp_path / "out", options)[0]
    pages = PdfReader(result.files["entry_bundle_pdf"]).pages
    texts = [page.extract_text() or "" for page in pages]

    assert len(pages) == 6
    assert "INVOICE-A" in texts[0] and "No.1-1/3" in texts[0]
    assert "No.1-2/3" in texts[1]
    assert "INSPECTION-A" in texts[2] and "No.1-3/3" in texts[2]
    assert "INVOICE-B" in texts[3] and "No.2-1/3" in texts[3]
    assert "No.2-2/3" in texts[4]
    assert "INSPECTION-B" in texts[5] and "No.2-3/3" in texts[5]


def test_entry_bundle_can_disable_overlay(tmp_path):
    from pypdf import PdfReader
    from reportlab.pdfgen import canvas
    from tidoc_print import PrintEntry, PrintOptions, build_print_package

    invoice = tmp_path / "invoice.pdf"
    doc = canvas.Canvas(str(invoice))
    doc.drawString(72, 760, "INVOICE")
    doc.save()
    entry = PrintEntry(entry_id="e1", title="北京理工大学", invoice_pdfs=[str(invoice)])
    options = PrintOptions(annotate=False, batch_note="批次甲", make_reimburse_doc=False, make_acceptance_doc=False)
    result = build_print_package([entry], tmp_path / "out", options)[0]
    text = PdfReader(result.files["entry_bundle_pdf"]).pages[0].extract_text() or ""
    assert "INVOICE" in text
    assert "No." not in text and "批次甲" not in text
