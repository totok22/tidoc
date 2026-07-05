"""报账说明 / 验收单 Word 生成。移植自 invoice2docx/engine.py。

按抬头强隔离：调用方对每个抬头分别生成，绝不把两个抬头混进同一份。
"""

from __future__ import annotations

import shutil
from copy import deepcopy
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

from docx import Document

from .docx_util import clear_data_rows, clone_row, ensure_template_shape, set_cell, set_paragraph_text
from .models import PersonProfile, PrintEntry, PrintItem

TEMPLATES_DIR = Path(__file__).parent / "templates"
DEFAULT_REIMBURSE_TEMPLATE = TEMPLATES_DIR / "报账说明模板.docx"
DEFAULT_ACCEPTANCE_TEMPLATE = TEMPLATES_DIR / "验收单模板.docx"

MONEY = Decimal("0.01")


def _money(v: Decimal) -> Decimal:
    return v.quantize(MONEY, rounding=ROUND_HALF_UP)


def _fmt_money(v: Decimal, currency: bool = False) -> str:
    return f"{'¥' if currency else ''}{_money(v):.2f}"


def _fmt_decimal(v: Decimal | None, places: int = 8) -> str:
    if v is None:
        return ""
    quant = Decimal("1." + "0" * places)
    text = f"{v.quantize(quant, rounding=ROUND_HALF_UP):f}"
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text


def _compact_item(entry: PrintEntry, storage_location: str) -> PrintItem:
    """把一张发票的多条明细压成验收单用的一行（移植 compact_acceptance_items）。"""
    first = entry.items[0] if entry.items else None
    suffix = "等" if len(entry.items) > 1 else ""
    quantity = sum((it.quantity for it in entry.items if it.quantity is not None), Decimal("0"))
    if quantity == Decimal("0"):
        quantity = Decimal("1")
    actual_base = (first.actual_name if first else "") or "发票物资"
    product_base = (first.product_name if first and first.product_name else actual_base) or actual_base
    return PrintItem(
        actual_name=actual_base + suffix,
        product_name=product_base + suffix,
        unit=first.unit if first else "",
        quantity=quantity,
        total=_money(entry.total),
        seller=entry.seller or (first.seller if first else ""),
        invoice_no=entry.invoice_no,
        storage_location=(first.storage_location if first and first.storage_location else storage_location),
    )


def generate_reimburse_doc(
    entries: list[PrintEntry],
    out_path: str | Path,
    document_date: str,
    profile: PersonProfile | None = None,
    template: str | Path = DEFAULT_REIMBURSE_TEMPLATE,
) -> Path:
    """生成报账说明 Word。移植自 update_reimburse_doc。"""
    template = Path(template)
    ensure_template_shape(template, 3, "报账说明")
    out_path = Path(out_path)
    shutil.copyfile(template, out_path)
    doc = Document(out_path)

    total = sum((e.total for e in entries), Decimal("0"))
    set_paragraph_text(doc.paragraphs[0], (
        f"机械与车辆学院申请支出{_fmt_money(total)} 元。方程式车队比赛物资采买。"
        f"人民币{_fmt_money(total)}元需打款至学生账户如下："
    ))
    p = profile or PersonProfile()
    if len(doc.paragraphs) > 2:
        set_paragraph_text(
            doc.paragraphs[1],
            f"学号：{p.student_id}   姓名：{p.person_name}   联系方式：{p.contact}",
        )
        set_paragraph_text(
            doc.paragraphs[2],
            f"开户行：{p.bank_name}       卡号：{p.bank_card}",
        )
    set_paragraph_text(doc.paragraphs[-1], document_date)

    table = doc.tables[0]
    template_row = deepcopy(table.rows[1])
    clear_data_rows(table, 1)
    for entry in entries:
        first = entry.items[0] if entry.items else PrintItem(actual_name="发票物资")
        suffix = "等" if len(entry.items) > 1 else ""
        row = clone_row(table, template_row)
        values = [
            (first.actual_name or "发票物资") + suffix,
            (first.actual_name or "发票物资") + suffix,
            _fmt_money(entry.total, currency=True),
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

    doc.save(out_path)
    return out_path


def generate_acceptance_doc(
    entries: list[PrintEntry],
    out_path: str | Path,
    document_date: str,
    storage_location: str = "工训楼",
    template: str | Path = DEFAULT_ACCEPTANCE_TEMPLATE,
) -> Path:
    """生成验收单 Word。移植自 update_acceptance_doc（每张发票压成一行）。"""
    template = Path(template)
    ensure_template_shape(template, 9, "验收单")
    out_path = Path(out_path)
    shutil.copyfile(template, out_path)
    doc = Document(out_path)

    set_paragraph_text(
        doc.paragraphs[1],
        f"单位   机械与车辆学院                                                                                       {document_date}",
    )
    table = doc.tables[0]
    template_row = deepcopy(table.rows[1])
    attachment_text = table.rows[1].cells[-1].text
    clear_data_rows(table, 1)

    for entry in entries:
        item = _compact_item(entry, storage_location)
        row = clone_row(table, template_row)
        values = [
            item.product_name or item.actual_name,
            item.unit,
            _fmt_decimal(item.quantity),
            _fmt_decimal(item.unit_price),
            _fmt_money(item.total),
            item.seller,
            item.invoice_no,
            item.storage_location or storage_location,
            attachment_text,
        ]
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

    doc.save(out_path)
    return out_path
