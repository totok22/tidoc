"""python-docx 底层辅助。移植自 invoice2docx/engine.py（clone_row 等）。"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document


def clone_row(table, template_row):
    """按模板行克隆一行追加到表尾，返回新行。"""
    new_tr = deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def clear_data_rows(table, keep_rows: int = 1):
    """清掉数据行，保留表头（默认保留 1 行）。"""
    for row in list(table.rows)[keep_rows:]:
        table._tbl.remove(row._tr)


def set_paragraph_text(paragraph, text: str):
    """替换段落文字，保留首个 run 的样式。"""
    text = str(text)
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def set_cell(cell, text: str):
    """写单元格，支持多行（按换行拆分）。"""
    lines = str(text).splitlines() or [""]
    for idx, paragraph in enumerate(cell.paragraphs):
        set_paragraph_text(paragraph, lines[idx] if idx < len(lines) else "")
    if len(lines) > len(cell.paragraphs):
        paragraph = cell.paragraphs[-1]
        for line in lines[len(cell.paragraphs):]:
            paragraph.add_run().add_break()
            paragraph.add_run(line)


def ensure_template_shape(template: Path, min_table_columns: int, label: str):
    """校验模板：至少要有一个表格、表头 + 样例数据行、足够列数。"""
    doc = Document(template)
    if not doc.tables:
        raise RuntimeError(f"{label}模板缺少表格：{template}")
    if len(doc.tables[0].rows) < 2:
        raise RuntimeError(f"{label}模板至少需要表头行和样例数据行：{template}")
    if len(doc.tables[0].rows[0].cells) < min_table_columns:
        raise RuntimeError(f"{label}模板表格列数不足：{template}")
